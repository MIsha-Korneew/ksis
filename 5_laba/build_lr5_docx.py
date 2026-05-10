# -*- coding: utf-8 -*-
"""Собрать ЛР_5_Корнеев.docx из текста отчёта (python-docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "ЛР_5_Корнеев.docx"


def set_body_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def p_center(doc: Document, text: str, bold: bool = False, size: int | None = None) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def p_line(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Normal")


def p_bold_lead(doc: Document, lead: str, rest: str) -> None:
    para = doc.add_paragraph()
    r1 = para.add_run(lead)
    r1.bold = True
    r1.font.name = "Times New Roman"
    r2 = para.add_run(rest)
    r2.font.name = "Times New Roman"


def add_table_curl(doc: Document) -> None:
    headers = ("Действие", "Пример команды", "Ожидаемый код (кратко)")
    rows = [
        (
            "Загрузить файл",
            "curl.exe -i -T hello.txt http://127.0.0.1:8765/demo/hello.txt",
            "201 или 200 при перезаписи",
        ),
        (
            "Список каталога",
            "curl.exe -i http://127.0.0.1:8765/demo/",
            "200 OK, тело — JSON",
        ),
        (
            "Скачать файл",
            "curl.exe -i http://127.0.0.1:8765/demo/hello.txt",
            "200 OK",
        ),
        (
            "Метаданные (без тела)",
            "curl.exe -I http://127.0.0.1:8765/demo/hello.txt",
            "200 OK, заголовки Content-Length, Last-Modified",
        ),
        (
            "Удалить",
            "curl.exe -i -X DELETE http://127.0.0.1:8765/demo/hello.txt",
            "204 No Content",
        ),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        for p in table.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(rows, 1):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell


def add_table_codes(doc: Document) -> None:
    headers = ("Код", "Расшифровка", "Где используется в работе")
    rows = [
        ("200", "OK", "Успешный GET / HEAD / повторный PUT"),
        ("201", "Created", "PUT создал новый файл"),
        ("204", "No Content", "Успешный DELETE"),
        ("403", "Forbidden", "Выход за пределы storage_data"),
        ("404", "Not Found", "Нет файла/каталога"),
        ("405", "Method Not Allowed", "HEAD для каталога"),
        ("409", "Conflict", "PUT, если путь — каталог"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        for p in table.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(rows, 1):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell


def main() -> None:
    doc = Document()
    set_body_style(doc)

    p_center(doc, "Министерство образования Республики Беларусь")
    p_center(
        doc,
        "Учреждение образования «Белорусский государственный университет\n"
        "информатики и радиоэлектроники»",
    )
    p_center(doc, "Факультет Компьютерного Проектирования")
    p_center(doc, "Кафедра инженерной психологии и эргономики")
    p_center(doc, "Дисциплина: Компьютерные системы и сети")
    doc.add_paragraph()

    p_center(doc, "ОТЧЁТ", bold=True, size=14)
    p_center(doc, "к лабораторной работе")
    p_center(doc, "на тему")
    p_center(
        doc,
        "«Файловое хранилище (REST API по HTTP)»",
        bold=True,
    )
    doc.add_paragraph()

    p_line(doc, "Выполнил:\t\t\tПроверила:")
    p_line(doc, "ст. гр. 410901 Болтак С.В.")
    p_line(doc, "Корнеев М.С.")
    doc.add_paragraph()

    p_center(doc, "Минск 2026")
    doc.add_paragraph()

    p_bold_lead(
        doc,
        "Цель работы: ",
        "закрепить прикладное использование протокола HTTP — методы запроса, заголовки, коды "
        "состояния; освоить вызовы HTTP API средствами curl (и при необходимости аналогов); "
        "реализовать удалённое файловое хранилище с REST-интерфейсом.",
    )

    p_bold_lead(
        doc,
        "Задание: ",
        "реализовать службу, в которой логический путь в URL задаёт положение файла в "
        "хранилище; поддержать PUT, GET (файл и список каталога JSON), HEAD (метаданные в "
        "заголовках), DELETE; корректные коды HTTP; проверка curl и браузера. "
        "Дополнительное задание (X-Copy-From) не выполнялось.",
    )

    doc.add_paragraph()
    p_bold_lead(doc, "Описание программной реализации", "")
    p_line(
        doc,
        "Реализация на Python 3, веб-фреймворк Flask. Файлы на диске в каталоге storage_data "
        "относительно storage_server.py (STORAGE_ROOT).",
    )
    p_line(
        doc,
        "Маршруты / и /<path> обрабатывают GET, PUT, HEAD, DELETE. Пути нормализуются; выход "
        "за пределы хранилища блокируется (403).",
    )
    p_line(
        doc,
        "PUT: запись тела в файл, 201 / 200, конфликт с каталогом — 409. "
        "GET: файл — send_file; каталог — JSON. HEAD: только файл — Content-Length, "
        "Last-Modified; каталог — 405. DELETE: файл или рекурсивно каталог — 204. "
        "После ответа в консоль: МЕТОД путь -> код (http_code_caption). "
        "Сервер: 0.0.0.0:8765.",
    )

    doc.add_paragraph()
    p_bold_lead(doc, "Проверка утилитой curl (в PowerShell: curl.exe)", "")
    add_table_curl(doc)
    p_line(
        doc,
        "Флаг -i выводит статусную строку ответа; -I — только заголовки.",
    )

    doc.add_paragraph()
    p_bold_lead(doc, "Анализ результатов работы программы", "")
    p_line(
        doc,
        "Рис. 1. Запуск сервера: каталог storage_data и адрес http://127.0.0.1:8765/.",
    )
    p_line(
        doc,
        "Рис. 2. Консоль: строки МЕТОД /путь -> код (расшифровка).",
    )
    p_line(
        doc,
        "Рис. 3. Вывод curl.exe -i со статусной строкой HTTP/1.x.",
    )
    p_line(
        doc,
        "Рис. 4. Браузер: JSON каталога или отображение файла.",
    )

    doc.add_paragraph()
    p_bold_lead(doc, "Таблица использованных кодов ответов (фрагмент)", "")
    add_table_codes(doc)

    doc.add_paragraph()
    p_bold_lead(doc, "Вывод по работе", "")
    p_line(
        doc,
        "Реализовано файловое хранилище с REST API по HTTP; добавлен вывод кодов с "
        "расшифровкой. Проверка curl.exe и браузером.",
    )
    p_line(
        doc,
        "Методичка (ЛР 5): https://sites.google.com/view/ksis-site/ "
        "(раздел «5. Файловое хранилище»).",
    )

    doc.add_paragraph()
    p_bold_lead(doc, "Список литературы: ", "документация Python/Flask; RFC 9110.")

    doc.save(OUT)
    print(f"Создан файл: {OUT}")


if __name__ == "__main__":
    main()
